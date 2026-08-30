from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "企业工作站使用说明.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
BORDER = "C8D1DC"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, latin=BODY_FONT):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_fixed_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    run._r.extend([begin, instr, separate, placeholder, end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def set_page_furniture(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("企业工作站  ·  业务闭环操作手册")
    set_run_font(r, size=9, color=MUTED, bold=True)

    # LibreOffice may map the default header to one side only even when Word's
    # odd/even setting is disabled. Populate the alternate part as a consistent
    # cross-renderer fallback while keeping Word on a unified running header.
    even_header = section.even_page_header
    p = even_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("企业工作站  ·  业务闭环操作手册")
    set_run_font(r, size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    add_page_field(p)

    even_footer = section.even_page_footer
    p = even_footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    add_page_field(p)


def patch_numbering_styles(doc):
    numbering = doc.part.numbering_part.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for style_name, fmt, text in (
        ("List Bullet", "bullet", "•"),
        ("List Number", "decimal", "%1."),
    ):
        style = doc.styles[style_name]
        ppr = style.element.get_or_add_pPr()
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is None:
            continue
        num_id_node = num_pr.find(qn("w:numId"))
        if num_id_node is None:
            continue
        num_id = num_id_node.get(qn("w:val"))
        num = numbering.find(f"w:num[@w:numId='{num_id}']", ns)
        if num is None:
            continue
        abs_id_node = num.find("w:abstractNumId", ns)
        abs_id = abs_id_node.get(qn("w:val"))
        abstract = numbering.find(f"w:abstractNum[@w:abstractNumId='{abs_id}']", ns)
        if abstract is None:
            continue
        lvl = abstract.find("w:lvl[@w:ilvl='0']", ns)
        if lvl is None:
            continue
        num_fmt = lvl.find("w:numFmt", ns)
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = lvl.find("w:lvlText", ns)
        lvl_text.set(qn("w:val"), text)
        suff = lvl.find("w:suff", ns)
        if suff is None:
            suff = OxmlElement("w:suff")
            lvl.append(suff)
        suff.set(qn("w:val"), "tab")
        ppr_lvl = lvl.find("w:pPr", ns)
        if ppr_lvl is None:
            ppr_lvl = OxmlElement("w:pPr")
            lvl.append(ppr_lvl)
        tabs = ppr_lvl.find("w:tabs", ns)
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            ppr_lvl.append(tabs)
        for old in list(tabs):
            tabs.remove(old)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = ppr_lvl.find("w:ind", ns)
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr_lvl.append(ind)
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    patch_numbering_styles(doc)


def add_para(doc, text, *, bold=False, color=INK, align=None, before=0, after=6, size=11, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    lvl_ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    lvl_ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    lvl_ppr.append(ind)
    lvl.append(lvl_ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph(style="List Number")
        ppr = p._p.get_or_add_pPr()
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            ppr.insert(0, num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=LIGHT_BLUE, size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_role_table(doc):
    headers = ["岗位类型", "岗位角色", "首次进入重点", "主要职责"]
    rows = [
        ["管理层", "决策人", "AI 决策调度台", "下达目标、确认拆解、处理风险、总验收"],
        ["业务负责人", "部门负责人", "负责人推进台", "承接目标、调度资源、验收成果、业务审批"],
        ["执行成员", "员工", "我的执行台", "执行任务、上报阻塞、上传成果、提交验收"],
        ["财务人员", "财务", "财务执行中心", "费用审批、薪资核对与发放"],
        ["人事人员", "人事", "组织人事", "人员档案、职级政策、薪资复核"],
    ]
    table = doc.add_table(rows=1, cols=4)
    set_fixed_table_geometry(table, [1150, 1450, 2500, 4260])
    set_table_borders(table)
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10, color=DARK_BLUE, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            set_cell_margins(cell)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(value)
            set_run_font(r, size=9.5, color=INK, bold=(i == 0))
    set_fixed_table_geometry(table, [1150, 1450, 2500, 4260])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("操作手册  ·  商业候选版")
    set_run_font(r, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("企业工作站")
    set_run_font(r, size=32, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("从目标到执行、验收、审批与复盘的真实业务闭环")
    set_run_font(r, size=14, color=DARK_BLUE)

    add_para(
        doc,
        "面向决策人、部门负责人、员工、财务与人事",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=54,
    )

    table = doc.add_table(rows=3, cols=2)
    set_fixed_table_geometry(table, [2700, 6660])
    set_table_borders(table, color=LIGHT_BLUE)
    values = [
        ("环境", "独立 Staging 验收后开放内部试用"),
        ("数据方式", "Supabase PostgreSQL + RLS + Storage"),
        ("版本日期", "2026 年 8 月 30 日"),
    ]
    for row, (label, value) in zip(table.rows, values):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for i, text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=DARK_BLUE if i == 0 else INK, bold=(i == 0))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    configure_styles(doc)
    for section in doc.sections:
        set_page_furniture(section)

    props = doc.core_properties
    props.title = "企业工作站使用说明"
    props.subject = "业务闭环、角色权限和操作验收手册"
    props.author = "企业工作站项目组"
    props.keywords = "企业工作站, AI任务拆解, 审批, 薪资, 项目管理"

    add_cover(doc)

    doc.add_heading("1. 一页上手", level=1)
    add_callout(
        doc,
        "系统目标",
        "领导只需说明目标、截止时间、预算和约束；系统将目标拆到部门和个人，并把执行、成果、验收、审批、通知和复盘连接成一条可追踪业务链。",
    )
    add_role_table(doc)

    doc.add_heading("完整业务闭环", level=2)
    add_numbers(doc, [
        "决策人输入战略目标。",
        "AI 生成部门目标、唯一负责人、个人任务、依赖关系和验收标准。",
        "决策人检查方案并确认下发。",
        "员工或负责人开始执行，更新进度并处理依赖。",
        "执行人上传真实成果文件并提交验收。",
        "指定验收人通过成果，或说明原因退回修改。",
        "阻塞、逾期和验收超时进入领导行动清单。",
        "全部任务完成后，领导总验收、复盘并归档。",
    ])

    heading = doc.add_heading("2. 登录与岗位权限", level=1)
    heading.paragraph_format.page_break_before = True
    add_para(doc, "通过企业统一登录入口进入系统。服务端根据组织、部门、职级、角色和权限码生成工作会话。")
    add_callout(doc, "访问控制", "导航只显示当前身份有权使用且已通过商业就绪门禁的模块；直接访问越权地址会由服务端拒绝。")
    add_para(doc, "退出后服务端会话失效，PWA 静态缓存同时清理。")
    add_callout(
        doc,
        "正式上线提醒",
        "内部试用前必须在独立 Staging 环境完成真实身份、数据库迁移、RLS、对象存储、桌面端和移动端验收。",
    )

    doc.add_heading("3. 决策人操作", level=1)
    doc.add_heading("3.1 下达命令", level=2)
    add_numbers(doc, [
        "进入“AI 决策调度台”。",
        "点击“发起新决策”。",
        "填写战略目标、硬性截止、预算上限和关键约束。",
        "点击 AI 拆解，检查部门目标、负责人、任务、依赖和验收标准。",
        "确认方案并下发；任务会同步进入专项项目、任务中心、负责人推进台和个人执行台。",
    ])
    doc.add_heading("3.2 管理风险和周报", level=2)
    add_para(doc, "“今日必须处理”按紧急程度汇总依赖异常、阻塞超时、任务逾期、验收超时和待领导审批。首页“本周执行摘要”统计完成率、推进中、待验收、阻塞、逾期、依赖风险和待审批数量；点击“复制周报”可直接生成文字摘要。")
    doc.add_heading("3.3 总验收与归档", level=2)
    add_para(doc, "只有全部任务完成且协同请求关闭后，才能提交总验收。总验收通过后完成归档，已验收成果将沉淀到业务闭环记录。")

    doc.add_heading("4. 部门负责人操作", level=1)
    add_numbers(doc, [
        "在“负责人推进台”承接部门目标。",
        "检查 AI 推荐人员，结合人员标签、工作负荷和能力调整安排。",
        "协调阻塞、前置依赖和跨部门支持。",
        "员工提交成果后，对照验收标准填写意见并通过或退回。",
        "在审批中心处理当前权限范围内的业务申请。",
    ])
    add_callout(doc, "职责分离", "负责人本人执行的任务不能自提自批，系统会自动交给决策人验收。")

    doc.add_heading("5. 员工操作", level=1)
    doc.add_heading("5.1 执行和交付", level=2)
    add_numbers(doc, [
        "在“我的执行台”查看任务、期限、验收标准和前置依赖。",
        "前置任务完成后点击“开始执行”。",
        "更新进度；遇到问题时填写阻塞原因并上报。",
        "需要预算、增员或培训时发起协同请求。",
        "上传成果文件，确认文件可下载后提交验收。",
    ])
    add_callout(doc, "强制规则", "没有成果文件的任务不能进入验收，执行人也不能绕过验收直接把任务改为完成。")
    doc.add_heading("5.2 个人业务", level=2)
    add_bullets(doc, [
        "审批中心：发起并跟踪报销、采购与合同流程。",
        "通知中心：处理任务、审批和业务异常提醒。",
        "薪资管理：员工只查看本人工资单。",
    ])

    doc.add_heading("6. 财务操作", level=1)
    add_numbers(doc, [
        "处理预算、采购等财务协同。",
        "核对工资政策、奖金、扣减和工资单范围。",
        "等待人事复核和领导批准。",
        "批准完成后执行最终发放并留存状态。",
    ])
    add_callout(doc, "薪资原则", "工资数据确认 → 财务核算 → 人事复核 → 授权批准 → 财务发放。所有状态来自服务端记录。")

    doc.add_heading("7. 人事操作", level=1)
    add_numbers(doc, [
        "处理增员、培训和人员支持。",
        "维护组织档案、部门、职级和汇报关系。",
        "配置部门与职级对应的薪资政策。",
        "复核员工范围、工资政策和计算结果。",
        "处理人员状态变更并保留审计记录。",
    ])

    heading = doc.add_heading("8. 项目管理", level=1)
    heading.paragraph_format.page_break_before = True
    add_bullets(doc, [
        "概览：目标、健康状态、成员和动态。",
        "里程碑：创建和查看关键节点。",
        "任务：创建、分配、评论和状态查看。",
        "甘特图：根据任务与里程碑日期自动生成。",
        "文件：上传真实文件内容，刷新后仍可下载。",
        "日报：成员提交今日完成、下一步、阻塞和支持需求。",
        "复盘：负责人维护结果、有效做法、经验教训、后续动作并收口风险。",
    ])
    add_callout(doc, "流程边界", "AI 决策创建的专项项目必须从执行工作台推进；项目详情不能绕过验收流程直接修改任务结果。")

    doc.add_heading("9. 活动、客户与分析", level=1)
    doc.add_heading("9.1 活动", level=2)
    add_para(doc, "活动与项目共用统一数据仓库。创建活动后会同步出现在活动推进和项目管理中，并自动生成策划、执行、推广、复盘四个阶段。")
    doc.add_heading("9.2 客户", level=2)
    add_para(doc, "客户管理支持新建客户、筛选、查看档案、记录跟进并推进客户阶段。客户资料由服务端保存，并按租户和权限隔离。")
    doc.add_heading("9.3 数据分析", level=2)
    add_para(doc, "数据分析向决策人和部门负责人展示执行趋势、交付日历、风险提醒和真实任务周报；部门负责人只查看本部门范围。")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.page_break_before = True
    spacer.paragraph_format.space_after = Pt(24)
    spacer.paragraph_format.keep_with_next = True
    doc.add_heading("10. 通知与行动清单", level=1)
    add_bullets(doc, [
        "顶部铃铛显示当前角色未读通知。",
        "点击通知进入处理页面并标为已读。",
        "通知中心支持逐条已读和全部已读。",
        "每个角色只收到职责范围内的新任务、阻塞、验收、逾期和审批提醒。",
    ])

    doc.add_heading("11. 文件与数据保存", level=1)
    add_para(doc, "正式业务记录保存在 Supabase PostgreSQL，文件保存在 Supabase Storage，并由 RLS、对象权限和服务端审计保护。浏览器不作为正式业务数据库。")
    doc.add_heading("正式部署建议", level=2)
    add_numbers(doc, [
        "配置 Supabase 数据库、Auth 和 Storage。",
        "建立数据库与对象存储备份策略。",
        "配置 HTTPS、企业域名、权限审计和日志保留周期。",
    ])

    doc.add_heading("12. 常见问题", level=1)
    questions = [
        ("为什么任务不能开始？", "检查前置依赖。所有前置任务完成后，按钮才会变为“开始执行”。"),
        ("为什么不能提交验收？", "任务必须处于执行中，并至少上传一个成果文件。"),
        ("为什么负责人不能验收自己的任务？", "系统执行职责分离，负责人自己的交付由决策人验收。"),
        ("为什么不能进行薪资核算？", "检查当前账号权限、有效薪资政策、工资周期和员工范围是否完整。"),
        ("为什么页面自动跳回首页？", "当前角色没有该页面权限，请从当前岗位导航进入允许模块。"),
        ("文件为什么不能下载？", "检查文件处理状态、对象存储权限、记录归属和当前会话；失败时保留错误码用于追踪。"),
    ]
    for index, (question, answer) in enumerate(questions):
        if index == 3:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.page_break_before = True
            spacer.paragraph_format.space_after = Pt(24)
            spacer.paragraph_format.keep_with_next = True
        doc.add_heading(question, level=3)
        add_para(doc, answer)

    doc.add_heading("13. 最终验收清单", level=1)
    add_numbers(doc, [
        "决策人下达目标并确认下发。",
        "执行人开始任务、上传成果并提交验收。",
        "负责人完成验收或退回修改。",
        "员工发起协同，财务或人事完成处理。",
        "员工发起报销，审批人完成通过、退回和再次提交。",
        "人事核对组织、职级和薪资政策。",
        "财务核算、人事复核、领导批准、财务发放。",
        "决策人查看通知和周报，完成总验收并归档。",
    ])
    add_callout(doc, "验收通过标准", "五类角色均只能看到职责所需内容；每个提交动作都有下一责任人；成果、审批、通知和汇总数据能在对应页面继续推进。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def audit_docx(path):
    with ZipFile(path) as zf:
        document = ET.fromstring(zf.read("word/document.xml"))
        styles = ET.fromstring(zf.read("word/styles.xml"))
        numbering = ET.fromstring(zf.read("word/numbering.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    sect = document.find(".//w:sectPr", ns)
    margins = sect.find("w:pgMar", ns)
    assertions = {
        "letter_page": sect.find("w:pgSz", ns).get(qn("w:w")) == "12240" and sect.find("w:pgSz", ns).get(qn("w:h")) == "15840",
        "one_inch_margins": all(margins.get(qn(f"w:{side}")) == "1440" for side in ("top", "right", "bottom", "left")),
        "heading_styles": all(styles.find(f".//w:style[@w:styleId='{style_id}']", ns) is not None for style_id in ("Heading1", "Heading2", "Heading3")),
        "real_numbering": len(numbering.findall(".//w:abstractNum", ns)) > 0,
        "fixed_tables": all(tbl.find("w:tblPr/w:tblLayout", ns) is not None for tbl in document.findall(".//w:tbl", ns)),
        "table_widths": all(tbl.find("w:tblPr/w:tblW", ns).get(qn("w:w")) == "9360" for tbl in document.findall(".//w:tbl", ns)),
        "table_indents": all(tbl.find("w:tblPr/w:tblInd", ns).get(qn("w:w")) == "120" for tbl in document.findall(".//w:tbl", ns)),
    }
    failed = [name for name, passed in assertions.items() if not passed]
    if failed:
        raise RuntimeError(f"DOCX preset audit failed: {', '.join(failed)}")
    print(f"Created and audited: {path}")


if __name__ == "__main__":
    output = build_document()
    audit_docx(output)
